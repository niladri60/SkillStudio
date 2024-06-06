#imports
import spacy
import pdfplumber
from docx import Document
import re
import pandas
import pytesseract
from PIL import Image
import docx
import fitz  # PyMuPDF
# Load NLP model
nlp = spacy.load('en_core_web_sm')

def extract_text_from_pdf(file_path):
    with pdfplumber.open(file_path) as pdf:
        text = ''
        for page in pdf.pages:
            text += page.extract_text()
    return text
def extract_text_from_docx(file_path):
    doc = Document(file_path)
    text = '\n'.join([para.text for para in doc.paragraphs])
    print("Text extraction from DOCX complete.")
    return text
def extract_text_with_font_size(docx_file):
    doc = docx.Document(docx_file)
    text_with_font_size = []

    for para in doc.paragraphs:
        for run in para.runs:
            if run.font.size:
                text_with_font_size.append((run.text, run.font.size.pt))
    
    return text_with_font_size
def extract_largest_text(text_with_font_size):
    if not text_with_font_size:
        return None
    
    largest_text = max(text_with_font_size, key=lambda x: x[1])
    return largest_text[0]

def extract_email(text):
    email_pattern = re.compile(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b')
    matches = email_pattern.findall(text)
    print(f"Email extraction complete: {matches[0] if matches else 'No email found'}")
    return matches[0] if matches else ''    
def extract_phone(text):
    # Improved phone number pattern to capture various formats
    phone_pattern = re.compile(r'\+?\d[\d -]{8,}\d')
    matches = phone_pattern.findall(text)
    print(f"Phone number extraction complete: {matches[0] if matches else 'No phone number found'}")
    return matches[0] if matches else ''
def extract_name(text):
    # Split the text by lines and iterate through each line
    for line in text.split('\n'):
        # Check if the line contains typical name format (First Name Last Name)
        if line.count(' ') >= 1:
            return tuple(line.strip().split(' ', 1))  # Split the line into first name and last name
    return None, None  # Return None if name format is not found
def parse_cv(file_path, file_type):
    if file_type == 'pdf':
        text = extract_text_from_pdf(file_path)
    elif file_type == 'docx':
        text_with_font_size = extract_text_with_font_size(file_path)
        largest_text = extract_largest_text(text_with_font_size)
        text = largest_text if largest_text else ""
    else:
        raise ValueError('Unsupported file type')
    
    email = extract_email(text)
    phone = extract_phone(text)
    name = extract_name(text)
    
    extracted_info = {
        'name': name,
        'email': email,
        'phone': phone,
        'education': [],
        'experience': [],
        'skills': [],
        'address': [],
        'certifications': [],
        'projects': []
    }
    
    sections = text.split('\n')
    current_section = None
    section_keywords = {
        'education': ['education', 'coursework'],
        'experience': ['experience', 'employment', 'work'],
        'skills': ['skills', 'technologies'],
        'address': ['address', 'location'],
        'certifications': ['certifications', 'awards', 'achievements'],
        'projects': ['projects', 'portfolio']
    }

    for line in sections:
        line_lower = line.lower()
        if any(keyword in line_lower for keyword in section_keywords['education']):
            current_section = 'education'
            continue
        elif any(keyword in line_lower for keyword in section_keywords['experience']):
            current_section = 'experience'
            continue
        elif any(keyword in line_lower for keyword in section_keywords['skills']):
            current_section = 'skills'
            continue
        elif any(keyword in line_lower for keyword in section_keywords['address']):
            current_section = 'address'
            continue
        elif any(keyword in line_lower for keyword in section_keywords['certifications']):
            current_section = 'certifications'
            continue
        elif any(keyword in line_lower for keyword in section_keywords['projects']):
            current_section = 'projects'
            continue

        if current_section and line.strip():
            extracted_info[current_section].append(line.strip())

    return extracted_info
